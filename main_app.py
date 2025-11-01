import streamlit as st
from docx import Document
import re
import io
import time

#1 Extract place holder

def extract_placeholders_with_context(file):
    doc = Document(file)
    placeholders = []
    for p in doc.paragraphs:
        matches = re.findall(r"\{\{(.*?)\}\}", p.text)
        for m in matches:
            placeholders.append({
                "key": m.strip(),
                "context": p.text.strip()
            })
    file.seek(0)
    return placeholders

def extract_blanks_with_context(file):
    doc = Document(file)
    placeholders = []
    for p in doc.paragraphs:
        text = p.text
        matches = re.finditer(r"_{3,}", text)
        for m in matches:
            before = text[:m.start()].strip().split()[-3:]
            label = "_".join(before) if before else "UnknownField"
            placeholders.append({
                "key": label,
                "context": text
            })
    file.seek(0)
    return placeholders


#2 Filling the placeholders
def fill_values(file, placeholders):
    if "step" not in st.session_state:
        st.session_state.step = 0
        st.session_state.answers = {}


    if file and st.session_state.step < len(placeholders):
        st.session_state.step = max(0, min(st.session_state.step, len(placeholders)-1))
        field = placeholders[st.session_state.step]
        key = field["key"]
        context = field["context"]
        st.markdown(f"**Context:** {context}")
        prev_val = st.session_state.answers.get(key, "")
        answer = st.text_input(f"Value for {key}:", value=prev_val)

        col1, col2 = st.columns(2)

        with col1:
            if st.button("Previous") and st.session_state.step > 0:
                #Save current answer before going back
                st.session_state.answers[key] = answer
                st.session_state.step -= 1
                st.rerun()

        with col2:
            if st.button("Next"):
                st.session_state.answers[key] = answer
                if st.session_state.step < len(placeholders) - 1:
                    st.session_state.step += 1
                    st.rerun()
                else:
                    st.session_state.step = len(placeholders)
                    st.rerun()


#3 Fill the document
def fill_docx(file, answers):
    file.seek(0)
    doc = Document(file)
    values = list(answers.values())
    val_index = 0
    for p in doc.paragraphs:
        while "___" in p.text and val_index < len(values):
            p.text = re.sub(r"_{3,}", values[val_index], p.text, count=1)
            val_index += 1

        for key, val in answers.items():
            p.text = p.text.replace(f"{{{{{key}}}}}", val)

            
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

#4 Generate section
def render_generate_section(file, placeholders):
    if file and st.session_state.step == len(placeholders):
        st.success("All fields filled!")
        if st.button("Generate Document"):
            out = fill_docx(file, st.session_state.answers)
            st.download_button(
                "Download Completed Document",
                data=out.getvalue(),
                file_name="completed.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Main app
def main():
    st.title("Legal Document Filler")
    if "step" not in st.session_state:
        st.session_state.step = 0
    if "answers" not in st.session_state:
        st.session_state.answers = {}

    uploaded_file = st.file_uploader("Upload a .docx file", type=["docx"])

        
    
    if not uploaded_file:
        st.info("Upload a .docx file with placeholders like {{ClientName}} to begin.")
        return
    with st.spinner("Extracting placeholders..."):
        time.sleep(1)
        placeholders = extract_placeholders_with_context(uploaded_file)
    if not placeholders:
        placeholders = extract_blanks_with_context(uploaded_file)
        if not placeholders: 
            st.warning("No placeholders found...")
            return
    fill_values(uploaded_file, placeholders)
    render_generate_section(uploaded_file, placeholders)

if __name__ == "__main__":
    main()
